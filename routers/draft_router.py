"""
智能编制助手API路由
提供制度文档管理、模板管理、草稿会话、AI生成大纲等接口
"""

import os
import json
import asyncio
import uuid
from datetime import datetime
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Header
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field
from db import get_db
from models.user import User
from models.draft import (
    DraftDocument, DocumentTemplate, DraftSession,
    DocumentStatus, DraftSessionStatus
)
from services.draft_service import DraftService
from routers.auth_router import get_current_user
from utils.logger import setup_logger

logger = setup_logger(__name__)

router = APIRouter()

UPLOAD_DIR = "uploads"
MATERIALS_DIR = os.path.join(UPLOAD_DIR, "materials")
ATTACHMENTS_DIR = os.path.join(UPLOAD_DIR, "attachments")
TEMPLATES_DIR = os.path.join(UPLOAD_DIR, "templates")

for d in [MATERIALS_DIR, ATTACHMENTS_DIR, TEMPLATES_DIR]:
    os.makedirs(d, exist_ok=True)


# ==================== 请求模型 ====================

class UserPermissionItem(BaseModel):
    user_id: int
    can_view: bool
    can_edit: bool


class SetPermissionsRequest(BaseModel):
    user_permissions: List[UserPermissionItem]


class CreateSessionRequest(BaseModel):
    template_id: int
    template_name: str
    document_type: str
    custom_name: str = ""
    creator_id: int


class SaveRelationsRequest(BaseModel):
    upper_documents: List[dict] = []
    lower_documents: List[dict] = []
    workflow_notes: str = ""


class SaveRequirementsRequest(BaseModel):
    requirements: str
    additional_notes: str = ""
    special_constraints: Optional[List[str]] = None


class GenerateOutlineRequest(BaseModel):
    document_title: str
    generation_options: Optional[dict] = None


class SaveDraftRequest(BaseModel):
    content: str
    content_format: str = "markdown"
    last_edited_chapter: str = ""
    auto_save: bool = False


class ExportOutlineWordRequest(BaseModel):
    include_toc: bool = True
    include_header_footer: bool = True
    header_text: str = ""
    footer_text: str = ""


class ExportFinalRequest(BaseModel):
    format: str = "docx"
    options: Optional[dict] = None


class SubmitReviewRequest(BaseModel):
    document_name: str
    document_type: str
    reviewers: List[int]
    review_deadline: Optional[str] = None
    priority: str = "normal"
    attachments: Optional[List[str]] = None
    submission_note: str = ""


class AiChatRequest(BaseModel):
    message: str
    context: Optional[dict] = None
    mode: str = "qna"


class ComplianceCheckRequest(BaseModel):
    check_scope: List[str]
    reference_documents: Optional[List[str]] = None


# ==================== 1. 主页面接口 ====================

@router.get("/statistics", summary="获取文档统计数据")
def get_statistics(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取智能编制助手主页面的统计卡片数据

    业务逻辑：统计当前用户的文档总数、起草中、已完成、待审查等
    """
    service = DraftService(db)
    return service.get_statistics(current_user.id)


@router.get("/list", summary="获取制度文档列表")
def get_document_list(
    keyword: str = "",
    status: Optional[str] = None,
    type: Optional[str] = None,
    skip: int = 0,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    分页获取用户的制度文档列表

    业务逻辑：查询用户创建或有权限查看的文档，支持搜索和筛选
    """
    service = DraftService(db)
    return service.get_document_list(
        user_id=current_user.id,
        keyword=keyword,
        status=status,
        doc_type=type,
        skip=skip,
        limit=limit
    )


@router.get("/{doc_id}/available-users", summary="获取可授权用户列表")
def get_available_users(
    doc_id: int,
    keyword: str = "",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取可用于文档授权的用户列表

    业务逻辑：查询所有活跃用户，并关联已有权限信息
    """
    service = DraftService(db)
    return service.get_available_users(doc_id, keyword)


@router.post("/{doc_id}/permissions", summary="设置文档用户权限")
def set_permissions(
    doc_id: int,
    request: SetPermissionsRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    批量设置指定文档的用户访问权限

    业务逻辑：遍历权限列表，已有则更新，否则创建
    """
    service = DraftService(db)
    return service.set_permissions(doc_id, [p.dict() for p in request.user_permissions])


# ==================== 2. 模板接口 ====================

@router.get("/template/list", summary="获取模板列表")
def get_template_list(
    type: Optional[str] = None,
    keyword: str = "",
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取系统预定义的制度模板列表

    业务逻辑：查询所有模板，支持按类型筛选和关键词搜索
    """
    service = DraftService(db)
    return service.get_template_list(
        template_type=type,
        keyword=keyword,
        skip=skip,
        limit=limit
    )


@router.get("/template/{template_id}", summary="获取模板详情")
def get_template_detail(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取指定模板的详细内容和结构信息

    业务逻辑：查询模板详情，包含章节结构、示例内容等
    """
    service = DraftService(db)
    template = service.get_template_by_id(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    return template.to_dict()


@router.post("/template/upload-custom", summary="上传自定义模板文件")
async def upload_custom_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    category: str = Form("自定义"),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    上传自定义模板文件

    业务逻辑：
    1. 保存文件到服务器
    2. 创建模板记录
    3. 尝试解析文档结构
    """
    allowed_ext = {".docx", ".doc", ".pdf"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=422, detail="不支持的文件格式")

    file_path = os.path.join(TEMPLATES_DIR, f"custom_{uuid.uuid4().hex[:8]}{ext}")
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    service = DraftService(db)
    template = service.create_custom_template(
        name=name,
        file_path=file_path,
        category=category,
        description=description,
        creator_id=current_user.id
    )

    return {
        "success": True,
        "template_id": template.id,
        "message": "模板上传成功",
        "file_path": file_path,
        "parsed_outline": {
            "chapters": template.content_structure or [],
            "word_count": 0
        }
    }


# ==================== 3. 关联制度接口 ====================

@router.get("/search-upper", summary="搜索上位制度/法律依据")
def search_upper_documents(
    keyword: str,
    document_type: Optional[str] = None,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    搜索可作为上位法或法律依据的相关制度文档

    业务逻辑：在已发布文档中搜索匹配关键词的文档
    """
    service = DraftService(db)
    return service.search_upper_documents(keyword, document_type, limit)


@router.get("/search-lower", summary="搜索下位制度/执行手册")
def search_lower_documents(
    keyword: str,
    parent_doc_type: Optional[str] = None,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    搜索可作为下位执行细则的相关文档

    业务逻辑：搜索匹配关键词的文档
    """
    service = DraftService(db)
    return service.search_lower_documents(keyword, parent_doc_type, limit)


@router.post("/{doc_id}/relations", summary="保存制度关联关系")
def save_document_relations(
    doc_id: str,
    request: SaveRelationsRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    保存当前制度与其他制度的上下位关系

    业务逻辑：清除旧关联，批量插入新关联
    """
    service = DraftService(db)
    return service.save_document_relations(
        doc_id=doc_id,
        upper_documents=request.upper_documents,
        lower_documents=request.lower_documents,
        workflow_notes=request.workflow_notes
    )


# ==================== 4. 上传资料接口 ====================

@router.post("/{draft_id}/upload-materials", summary="上传参考材料文件")
async def upload_materials(
    draft_id: str,
    files: List[UploadFile] = File(...),
    material_type: str = Form("reference"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    批量上传制度编制所需的参考文档

    业务逻辑：
    1. 保存文件到服务器
    2. 创建材料记录
    3. 更新会话状态
    """
    service = DraftService(db)
    session = service.get_session(draft_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    allowed_ext = {".doc", ".docx", ".xls", ".xlsx", ".pdf", ".txt", ".md"}
    uploaded_files = []

    for file in files:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_ext:
            continue

        file_id = f"file_{uuid.uuid4().hex[:8]}"
        save_path = os.path.join(MATERIALS_DIR, f"{file_id}{ext}")
        content = await file.read()

        if len(content) > 10 * 1024 * 1024:
            continue

        with open(save_path, "wb") as f:
            f.write(content)

        uploaded_files.append({
            "file_name": file.filename,
            "file_path": save_path,
            "file_size": len(content),
            "file_type": ext.replace(".", "")
        })

    result = service.upload_materials(draft_id, uploaded_files, material_type)
    return result


@router.delete("/{draft_id}/material/{file_id}", summary="删除已上传的参考材料")
def delete_material(
    draft_id: str,
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    删除指定的参考材料文件

    业务逻辑：删除文件记录和物理文件
    """
    service = DraftService(db)
    if not service.delete_material(draft_id, file_id):
        raise HTTPException(status_code=404, detail="文件不存在")
    return {
        "success": True,
        "message": "文件删除成功",
        "deleted_file_id": file_id
    }


@router.post("/{draft_id}/requirements", summary="保存需求和补充信息")
def save_requirements(
    draft_id: str,
    request: SaveRequirementsRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    保存用户对制度的具体要求和补充备注

    业务逻辑：更新会话的需求字段和特殊约束条件
    """
    service = DraftService(db)
    return service.save_requirements(
        session_id=draft_id,
        requirements=request.requirements,
        additional_notes=request.additional_notes,
        special_constraints=request.special_constraints
    )


# ==================== 5. 生成大纲接口 ====================

@router.post("/create-session", summary="创建草稿会话")
def create_session(
    request: CreateSessionRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    初始化一个新的制度草稿会话

    业务逻辑：
    1. 生成唯一会话ID
    2. 创建会话和关联文档记录
    3. 增加模板使用次数
    """
    service = DraftService(db)
    session = service.create_session(
        template_id=request.template_id,
        template_name=request.template_name,
        document_type=request.document_type,
        creator_id=current_user.id,
        custom_name=request.custom_name
    )
    return {
        "success": True,
        "draft_session": session.to_dict()
    }


@router.post("/{session_id}/generate-outline", summary="AI生成制度大纲(SSE)")
async def generate_outline(
    session_id: str,
    request: GenerateOutlineRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    基于收集的信息调用AI生成制度大纲，SSE流式返回

    业务逻辑：
    1. 收集会话中的模板、关联文档、参考资料、需求信息
    2. 构建AI提示词
    3. 调用AI模型流式生成大纲
    4. 保存生成结果
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    async def event_generator():
        try:
            yield f"data: {json.dumps({'event': 'generation_started', 'message': '开始生成大纲...', 'timestamp': datetime.now().isoformat()}, ensure_ascii=False)}\n\n"
            service.update_session_status(session_id, DraftSessionStatus.GENERATING.value)

            context_parts = []
            context_parts.append(f"制度名称：{request.document_title}")
            context_parts.append(f"制度类型：{session.document_type}")

            if session.template_id:
                template = service.get_template_by_id(session.template_id)
                if template and template.content_structure:
                    context_parts.append(f"参考模板结构：{json.dumps(template.content_structure, ensure_ascii=False)}")
                if template and template.sample_content:
                    context_parts.append(f"模板示例内容：\n{template.sample_content[:2000]}")

            relations = service.get_document_relations(session_id)
            if relations["upper_documents"]:
                upper_names = [d.get("related_doc_name", "") for d in relations["upper_documents"]]
                context_parts.append(f"上位制度/法律依据：{', '.join(upper_names)}")

            materials = service.get_materials(session_id)
            if materials:
                mat_names = [m.get("file_name", "") for m in materials]
                context_parts.append(f"参考资料：{', '.join(mat_names)}")

            if session.requirements:
                context_parts.append(f"需求说明：{session.requirements}")
            if session.special_constraints:
                context_parts.append(f"特殊约束：{', '.join(session.special_constraints)}")

            gen_opts = request.generation_options or {}
            detail_level = gen_opts.get("detail_level", "standard")
            style = gen_opts.get("style", "formal")
            include_examples = gen_opts.get("include_examples", True)

            context_parts.append(f"详细程度：{detail_level}")
            context_parts.append(f"文风风格：{style}")
            context_parts.append(f"是否包含示例：{'是' if include_examples else '否'}")

            prompt = f"""请根据以下信息生成一份制度文档大纲：

{chr(10).join(context_parts)}

请生成完整的制度大纲，包含章节标题和条款要点。使用Markdown格式，章节用##标记，条款用###标记。"""

            yield f"data: {json.dumps({'event': 'progress_update', 'progress': 20, 'current_chapter': '正在准备生成...', 'is_end': False}, ensure_ascii=False)}\n\n"

            from config import settings
            if settings.MODEL_PROVIDER == "alibaba":
                from langchain_alibaba import ChatTongyi
                llm = ChatTongyi(
                    model=settings.ALIBABA_MODEL,
                    api_key=settings.ALIBABA_API_KEY,
                    streaming=True
                )
            else:
                from langchain_ollama import ChatOllama
                llm = ChatOllama(
                    base_url=settings.OLLAMA_BASE_URL,
                    model=settings.OLLAMA_MODEL,
                    streaming=True
                )

            from langchain_core.messages import HumanMessage

            full_content = ""
            progress = 30

            async for chunk in llm.astream([HumanMessage(content=prompt)]):
                content = chunk.content
                full_content += content

                yield f"data: {json.dumps({'event': 'outline_chunk', 'content': content, 'is_end': False}, ensure_ascii=False)}\n\n"

                if len(full_content) % 500 == 0:
                    progress = min(90, progress + 10)
                    yield f"data: {json.dumps({'event': 'progress_update', 'progress': progress, 'is_end': False}, ensure_ascii=False)}\n\n"

            outline_id = f"outline_{uuid.uuid4().hex[:8]}"
            service.save_outline(session_id, full_content, outline_id)

            yield f"data: {json.dumps({'event': 'generation_completed', 'message': '大纲生成成功！', 'total_tokens': len(full_content), 'generation_time_ms': 0, 'outline_id': outline_id, 'is_end': True}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"生成大纲失败: {str(e)}", exc_info=True)
            service.update_session_status(session_id, DraftSessionStatus.GENERATION_FAILED.value)
            yield f"data: {json.dumps({'event': 'generation_error', 'error_code': 'GENERATION_FAILED', 'error_message': str(e), 'suggestion': '请稍后重试'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/{session_id}/regenerate-outline", summary="重新生成大纲(SSE)")
async def regenerate_outline(
    session_id: str,
    request: GenerateOutlineRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    重新触发AI生成大纲，清除之前的内容

    业务逻辑：与generate-outline相同，但会清除旧大纲
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    session.outline_content = None
    session.outline_id = None
    db.commit()

    return await generate_outline(session_id, request, db, current_user)


@router.get("/{session_id}/outline", summary="获取已生成的大纲内容")
def get_outline(
    session_id: str,
    format: str = "markdown",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取之前生成的完整大纲内容

    业务逻辑：查询会话的大纲内容，支持markdown/json格式
    """
    service = DraftService(db)
    result = service.get_outline(session_id, format)
    if not result:
        raise HTTPException(status_code=404, detail="大纲不存在")
    return result


@router.post("/{session_id}/export-outline-word", summary="导出大纲为Word文档")
def export_outline_word(
    session_id: str,
    request: ExportOutlineWordRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    将大纲导出为Word文档

    业务逻辑：
    1. 获取大纲内容
    2. 生成Word文件
    3. 返回下载链接
    """
    service = DraftService(db)
    outline = service.get_outline(session_id, "markdown")
    if not outline or not outline.get("content"):
        raise HTTPException(status_code=404, detail="大纲内容不存在")

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋_GB2312'
        font.size = Pt(14)

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        if request.include_header_footer and request.header_text:
            header = sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = request.header_text.replace("{document_title}", outline.get("document_title", ""))

        content = outline.get("content", "")
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                p = doc.add_heading(line.replace("## ", ""), level=1)
            elif line.startswith("### "):
                p = doc.add_heading(line.replace("### ", ""), level=2)
            elif line.startswith("# "):
                p = doc.add_heading(line.replace("# ", ""), level=0)
            else:
                doc.add_paragraph(line)

        file_name = f"{outline.get('document_title', '大纲')}_草案大纲.docx"
        file_path = os.path.join(UPLOAD_DIR, "temp", f"outline_{uuid.uuid4().hex[:8]}.docx")
        doc.save(file_path)

        return {
            "success": True,
            "download_url": f"/api/file/download?path={file_path}&filename={file_name}",
            "file_name": file_name,
            "file_size": os.path.getsize(file_path),
            "expires_at": (datetime.now().replace(hour=23, minute=59)).isoformat()
        }
    except Exception as e:
        logger.error(f"导出Word失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


# ==================== 6. 编辑完善接口 ====================

@router.put("/{session_id}/save-draft", summary="保存制度草稿")
def save_draft(
    session_id: str,
    request: SaveDraftRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    保存当前编辑中的制度内容

    业务逻辑：更新会话和文档内容，手动保存时增加版本号
    """
    service = DraftService(db)
    return service.save_draft(
        session_id=session_id,
        content=request.content,
        content_format=request.content_format,
        last_edited_chapter=request.last_edited_chapter,
        auto_save=request.auto_save
    )


@router.post("/{session_id}/export-final", summary="导出最终版本文档")
def export_final(
    session_id: str,
    request: ExportFinalRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    将编辑完成的制度导出为最终格式文档

    业务逻辑：
    1. 获取草稿内容
    2. 根据选项生成正式文档
    3. 返回下载链接
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    content = session.draft_content or session.outline_content or ""
    if not content:
        raise HTTPException(status_code=400, detail="没有可导出的内容")

    options = request.options or {}

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = options.get("font_family", "仿宋_GB2312")
        font.size = Pt(options.get("body_font_size", 14))

        sections = doc.sections
        margins = options.get("margins", {})
        for section in sections:
            section.top_margin = Cm(margins.get("top", 2.54))
            section.bottom_margin = Cm(margins.get("bottom", 2.54))
            section.left_margin = Cm(margins.get("left", 3.17))
            section.right_margin = Cm(margins.get("right", 3.17))

        if options.get("include_cover_page", True):
            cover_info = options.get("cover_info", {})
            title = cover_info.get("title", session.custom_name or session.template_name)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(options.get("title_font_size", 22))
            run.bold = True

            for key in ["document_number", "version", "department", "date"]:
                val = cover_info.get(key, "")
                if val:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(val)

            doc.add_page_break()

        for line in content.split("\n"):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if line_stripped.startswith("# "):
                doc.add_heading(line_stripped.replace("# ", ""), level=0)
            elif line_stripped.startswith("## "):
                doc.add_heading(line_stripped.replace("## ", ""), level=1)
            elif line_stripped.startswith("### "):
                doc.add_heading(line_stripped.replace("### ", ""), level=2)
            elif line_stripped.startswith("#### "):
                doc.add_heading(line_stripped.replace("#### ", ""), level=3)
            else:
                doc.add_paragraph(line_stripped)

        if options.get("include_approval_section", True):
            doc.add_page_break()
            doc.add_heading("审批签批栏", level=1)
            for label in ["起草人", "部门负责人", "分管领导", "总经理"]:
                p = doc.add_paragraph()
                p.add_run(f"{label}：________________    日期：________________")

        doc_title = session.custom_name or session.template_name
        version_str = cover_info.get("version", f"V{session.version}")
        file_name = f"{doc_title}_{version_str}.docx"
        file_path = os.path.join(UPLOAD_DIR, "temp", f"final_{uuid.uuid4().hex[:8]}.docx")
        doc.save(file_path)

        return {
            "success": True,
            "download_url": f"/api/file/download?path={file_path}&filename={file_name}",
            "file_name": file_name,
            "file_size": os.path.getsize(file_path),
            "page_count": len(doc.sections),
            "expires_at": (datetime.now().replace(hour=23, minute=59)).isoformat()
        }
    except Exception as e:
        logger.error(f"导出最终文档失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/{session_id}/submit-review", summary="提交制度审核")
def submit_review(
    session_id: str,
    request: SubmitReviewRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    将编辑完成的制度提交进入审核流程

    业务逻辑：
    1. 更新会话和文档状态
    2. 设置审核人权限
    3. 返回审核任务信息
    """
    service = DraftService(db)
    return service.submit_review(
        session_id=session_id,
        document_name=request.document_name,
        document_type=request.document_type,
        reviewers=request.reviewers,
        review_deadline=request.review_deadline,
        priority=request.priority,
        submission_note=request.submission_note
    )


@router.post("/{session_id}/ai-chat", summary="AI助手对话(SSE)")
async def ai_chat(
    session_id: str,
    request: AiChatRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    在编辑页面与AI助手对话，SSE流式返回

    业务逻辑：
    1. 构建包含当前文档上下文的提示词
    2. 调用AI模型流式回复
    3. 支持问答、修改建议、智能扩充三种模式
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    async def event_generator():
        try:
            yield f"data: {json.dumps({'event': 'thinking', 'content': '', 'is_end': False}, ensure_ascii=False)}\n\n"

            context = request.context or {}
            mode_prompts = {
                "qna": "请回答以下关于制度文档的问题：",
                "modify": "请根据以下修改建议，提供具体的修改方案：",
                "expand": "请根据以下要求，扩充相关内容："
            }
            mode = request.mode or "qna"
            mode_prompt = mode_prompts.get(mode, mode_prompts["qna"])

            system_parts = [f"你是制度编制AI助手，正在协助用户编制《{session.custom_name or session.template_name}》制度文档。"]

            if session.draft_content:
                current_chapter = context.get("current_chapter", "")
                if current_chapter:
                    system_parts.append(f"当前编辑章节：{current_chapter}")

                selected_text = context.get("selected_text", "")
                if selected_text:
                    system_parts.append(f"用户选中的文本：{selected_text}")

            conversation_history = context.get("conversation_history", [])
            history_msgs = []
            for msg in conversation_history[-10:]:
                role = msg.get("role", "user")
                content = msg.get("content", "")
                if role == "assistant":
                    from langchain_core.messages import AIMessage
                    history_msgs.append(AIMessage(content=content))
                else:
                    from langchain_core.messages import HumanMessage
                    history_msgs.append(HumanMessage(content=content))

            full_prompt = f"{mode_prompt}\n\n{request.message}"
            from langchain_core.messages import HumanMessage, SystemMessage

            messages = [SystemMessage(content="\n".join(system_parts))] + history_msgs + [HumanMessage(content=full_prompt)]

            from config import settings
            if settings.MODEL_PROVIDER == "alibaba":
                from langchain_alibaba import ChatTongyi
                llm = ChatTongyi(
                    model=settings.ALIBABA_MODEL,
                    api_key=settings.ALIBABA_API_KEY,
                    streaming=True
                )
            else:
                from langchain_ollama import ChatOllama
                llm = ChatOllama(
                    base_url=settings.OLLAMA_BASE_URL,
                    model=settings.OLLAMA_MODEL,
                    streaming=True
                )

            tokens_used = 0
            async for chunk in llm.astream(messages):
                content = chunk.content
                tokens_used += len(content)

                if mode == "modify" and "建议" in content:
                    yield f"data: {json.dumps({'event': 'suggested_modification', 'original_text': context.get('selected_text', ''), 'suggested_text': content, 'explanation': ''}, ensure_ascii=False)}\n\n"
                else:
                    yield f"data: {json.dumps({'event': 'content_chunk', 'content': content, 'is_end': False}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'event': 'completed', 'message': '回答完成', 'tokens_used': tokens_used, 'is_end': True}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"AI对话失败: {str(e)}", exc_info=True)
            yield f"data: {json.dumps({'event': 'error', 'error_code': 'AI_CHAT_FAILED', 'error_message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/{session_id}/compliance-check", summary="合规检查")
async def compliance_check(
    session_id: str,
    request: ComplianceCheckRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    对当前制度内容进行自动化合规检查

    业务逻辑：
    1. 获取草稿内容
    2. 构建检查提示词
    3. 调用AI进行检查
    4. 解析并保存检查结果
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    content = session.draft_content or session.outline_content or ""
    if not content:
        raise HTTPException(status_code=400, detail="没有可检查的内容")

    scope_map = {
        "legal_compliance": "法律合规性检查（与上位法冲突检测）",
        "format_standard": "格式规范性检查（公文格式、条款编号等）",
        "internal_consistency": "内部一致性检查（前后矛盾、重复定义等）",
        "terminology": "术语规范性检查（专业术语使用是否准确）"
    }

    scope_descriptions = [scope_map.get(s, s) for s in request.check_scope]

    prompt = f"""请对以下制度文档内容进行合规检查，检查范围包括：{', '.join(scope_descriptions)}

制度内容：
{content[:5000]}

请以JSON格式返回检查结果，格式如下：
{{
    "overall_status": "pass/warning/fail",
    "summary": {{
        "total_issues": 数字,
        "critical": 数字,
        "warning": 数字,
        "info": 数字
    }},
    "issues": [
        {{
            "severity": "critical/warning/info",
            "category": "检查类别",
            "location": "问题位置",
            "description": "问题描述",
            "suggestion": "修改建议",
            "auto_fixable": false
        }}
    ],
    "passed_checks": [
        {{"check_type": "检查类型", "message": "通过信息"}}
    ]
}}

只返回JSON，不要其他内容。"""

    try:
        from config import settings
        if settings.MODEL_PROVIDER == "alibaba":
            from langchain_alibaba import ChatTongyi
            llm = ChatTongyi(model=settings.ALIBABA_MODEL, api_key=settings.ALIBABA_API_KEY)
        else:
            from langchain_ollama import ChatOllama
            llm = ChatOllama(base_url=settings.OLLAMA_BASE_URL, model=settings.OLLAMA_MODEL)

        from langchain_core.messages import HumanMessage
        response = await llm.ainvoke([HumanMessage(content=prompt)])
        result_text = response.content

        try:
            json_str = result_text
            if "```json" in json_str:
                json_str = json_str.split("```json")[1].split("```")[0]
            elif "```" in json_str:
                json_str = json_str.split("```")[1].split("```")[0]
            check_result = json.loads(json_str.strip())
        except json.JSONDecodeError:
            check_result = {
                "overall_status": "warning",
                "summary": {"total_issues": 0, "critical": 0, "warning": 0, "info": 0},
                "issues": [],
                "passed_checks": [{"check_type": "all", "message": "AI返回结果解析失败，请重新检查"}]
            }

        check = service.save_compliance_check(
            session_id=session_id,
            check_scope=request.check_scope,
            reference_documents=request.reference_documents,
            overall_status=check_result.get("overall_status", "warning"),
            summary=check_result.get("summary", {}),
            issues=check_result.get("issues", []),
            passed_checks=check_result.get("passed_checks", [])
        )

        return check.to_dict()

    except Exception as e:
        logger.error(f"合规检查失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"合规检查失败: {str(e)}")


@router.get("/{session_id}/references", summary="引用查看")
def get_references(
    session_id: str,
    reference_id: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取制度内容中引用的外部文档详细信息

    业务逻辑：查询关联的上位文档和参考资料
    """
    service = DraftService(db)
    return service.get_references(session_id, reference_id)


# ==================== 7. 附件上传接口 ====================

@router.post("/{session_id}/upload-attachment", summary="编辑器附件上传")
async def upload_attachment(
    session_id: str,
    file: UploadFile = File(...),
    attachment_type: str = Form("other"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    在编辑页面上传附件（图片、表格等）

    业务逻辑：
    1. 保存文件
    2. 创建附件记录
    3. 返回Markdown引用语法
    """
    service = DraftService(db)
    session = service.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="草稿会话不存在")

    ext = os.path.splitext(file.filename)[1].lower()
    image_ext = {".jpg", ".jpeg", ".png", ".gif", ".svg"}
    table_ext = {".xls", ".xlsx", ".csv"}
    other_ext = {".pdf"}

    if ext not in image_ext | table_ext | other_ext:
        raise HTTPException(status_code=422, detail="不支持的文件格式")

    if ext in image_ext and file.size and file.size > 5 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="图片文件不能超过5MB")
    if ext in table_ext and file.size and file.size > 2 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="表格文件不能超过2MB")

    if not attachment_type or attachment_type == "other":
        if ext in image_ext:
            attachment_type = "image"
        elif ext in table_ext:
            attachment_type = "table"
        else:
            attachment_type = "other"

    att_id = f"att_{uuid.uuid4().hex[:8]}"
    save_path = os.path.join(ATTACHMENTS_DIR, f"{att_id}{ext}")
    content = await file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    file_info = {
        "file_name": file.filename,
        "file_path": save_path,
        "file_size": len(content),
        "file_type": ext.replace(".", "")
    }

    attachment = service.upload_attachment(session_id, file_info, attachment_type)

    return {
        "success": True,
        "attachment": attachment.to_dict()
    }
