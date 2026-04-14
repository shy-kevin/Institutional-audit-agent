"""
审查相关API路由
"""
import json
import asyncio
import os
import re
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from db import get_db
from services.audit_service import AuditService
from models.audit import (
    AuditTask, AuditConfig, AuditChecklist, AuditResult, AuditIssue, AuditTrail,
    VersionCompareTask
)
from models.user import User
from agent.audit_agent import create_audit_agent
from utils.logger import setup_logger
from routers.auth_router import get_current_user

logger = setup_logger(__name__)


router = APIRouter(tags=["审查管理"])


# ==================== 请求模型 ====================

class CreateTaskRequest(BaseModel):
    document_path: str
    document_name: str
    audit_type: str = "draft"
    config_id: Optional[int] = None


class BatchCreateTaskRequest(BaseModel):
    documents: List[dict]
    audit_type: str = "draft"
    config_id: Optional[int] = None
    auto_start: bool = True


class CreateConfigRequest(BaseModel):
    name: str
    audit_dimensions: Optional[List[str]] = None
    focus_keywords: Optional[List[str]] = None
    checklist_ids: Optional[List[int]] = None
    is_default: bool = False


class StartTaskRequest(BaseModel):
    config_id: Optional[int] = None


class UpdateIssueStatusRequest(BaseModel):
    status: str
    suggestion: Optional[str] = None
    reject_reason: Optional[str] = None


class BatchUpdateIssuesRequest(BaseModel):
    issue_ids: List[int]
    status: str


class ConfirmResultRequest(BaseModel):
    comment: Optional[str] = None


class RejectResultRequest(BaseModel):
    reason: str


class ExportReportRequest(BaseModel):
    format: str = "word"


class CreateVersionCompareRequest(BaseModel):
    old_document_path: str
    new_document_path: str
    config_id: Optional[int] = None


# ==================== 响应转换函数 ====================

def task_to_dict(task: AuditTask) -> dict:
    return {
        "id": task.id,
        "document_name": task.document_name,
        "document_path": task.document_path,
        "status": task.status.value if task.status else "pending",
        "audit_type": task.audit_type.value if task.audit_type else "draft",
        "progress": task.progress or 0,
        "created_at": task.created_at.isoformat() if task.created_at else None,
        "updated_at": task.updated_at.isoformat() if task.updated_at else None,
        "completed_at": task.completed_at.isoformat() if task.completed_at else None
    }


def result_to_dict(result: AuditResult) -> dict:
    return {
        "id": result.id,
        "task_id": result.task_id,
        "document_name": result.document_name,
        "risk_level": result.risk_level.value if result.risk_level else None,
        "total_issues": result.total_issues or 0,
        "compliance_issues": result.compliance_issues or 0,
        "consistency_issues": result.consistency_issues or 0,
        "format_issues": result.format_issues or 0,
        "status": result.status.value if result.status else "pending_review",
        "created_at": result.created_at.isoformat() if result.created_at else None,
        "updated_at": result.updated_at.isoformat() if result.updated_at else None
    }


def issue_to_dict(issue: AuditIssue) -> dict:
    return {
        "id": issue.id,
        "result_id": issue.result_id,
        "issue_type": issue.issue_type.value if issue.issue_type else None,
        "severity": issue.severity.value if issue.severity else None,
        "location": issue.location,
        "original_text": issue.original_text,
        "issue_description": issue.issue_description,
        "legal_basis": issue.legal_basis,
        "suggestion": issue.suggestion,
        "status": issue.status.value if issue.status else "pending",
        "reject_reason": issue.reject_reason,
        "created_at": issue.created_at.isoformat() if issue.created_at else None
    }


def trail_to_dict(trail: AuditTrail) -> dict:
    return {
        "id": trail.id,
        "task_id": trail.task_id,
        "action": trail.action,
        "actor": trail.actor,
        "details": trail.details,
        "timestamp": trail.created_at.isoformat() if trail.created_at else None
    }


def config_to_dict(config: AuditConfig) -> dict:
    return {
        "id": config.id,
        "name": config.name,
        "audit_dimensions": config.audit_dimensions or [],
        "focus_keywords": config.focus_keywords or [],
        "checklist_ids": config.checklist_ids or [],
        "is_default": config.is_default or False,
        "created_at": config.created_at.isoformat() if config.created_at else None,
        "updated_at": config.updated_at.isoformat() if config.updated_at else None
    }


def checklist_to_dict(checklist: AuditChecklist) -> dict:
    return {
        "id": checklist.id,
        "name": checklist.name,
        "category": checklist.category.value if checklist.category else "general",
        "items": checklist.items or [],
        "is_active": checklist.is_active or False,
        "created_at": checklist.created_at.isoformat() if checklist.created_at else None,
        "updated_at": checklist.updated_at.isoformat() if checklist.updated_at else None
    }


def version_compare_to_dict(task: VersionCompareTask) -> dict:
    return {
        "id": task.id,
        "old_document_name": task.old_document_name,
        "new_document_name": task.new_document_name,
        "additions": task.additions or [],
        "deletions": task.deletions or [],
        "modifications": task.modifications or [],
        "created_at": task.created_at.isoformat() if task.created_at else None
    }


# ==================== 工作台接口 ====================

@router.get("/statistics")
def get_statistics(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取审查统计数据"""
    service = AuditService(db)
    return service.get_statistics(user_id=current_user.id)


@router.get("/tasks")
def get_tasks(
    limit: int = Query(10, ge=1, le=100),
    status: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取审查任务列表"""
    service = AuditService(db)
    result = service.get_tasks(limit=limit, status=status, user_id=current_user.id)
    return {
        "total": result["total"],
        "items": [task_to_dict(t) for t in result["items"]]
    }


@router.get("/history")
def get_history(
    date_range: Optional[str] = Query(None),
    audit_type: Optional[str] = Query(None),
    risk_level: Optional[str] = Query(None),
    keyword: Optional[str] = Query(None),
    limit: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取审查历史记录"""
    service = AuditService(db)
    result = service.get_history(
        date_range=date_range,
        audit_type=audit_type,
        risk_level=risk_level,
        keyword=keyword,
        limit=limit,
        user_id=current_user.id
    )
    return {
        "total": result["total"],
        "items": [result_to_dict(r) for r in result["items"]]
    }


# ==================== 任务管理接口 ====================

@router.post("/task/create")
def create_task(
    request: CreateTaskRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """创建审查任务"""
    service = AuditService(db)
    task = service.create_task(
        document_path=request.document_path,
        document_name=request.document_name,
        audit_type=request.audit_type,
        config_id=request.config_id,
        user_id=current_user.id
    )
    return task_to_dict(task)


@router.post("/task/batch-create")
def batch_create_tasks(
    request: BatchCreateTaskRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """批量创建审查任务并自动开始审查"""
    service = AuditService(db)
    tasks = []
    
    for doc in request.documents:
        task = service.create_task(
            document_path=doc.get("document_path", ""),
            document_name=doc.get("document_name", "未命名文档"),
            audit_type=request.audit_type,
            config_id=request.config_id,
            user_id=current_user.id
        )
        tasks.append(task_to_dict(task))
    
    if request.auto_start:
        for task in tasks:
            service.start_task(task["id"], request.config_id)
    
    return {
        "success": True,
        "total": len(tasks),
        "auto_started": request.auto_start,
        "tasks": tasks
    }


@router.get("/task/{task_id}")
def get_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取审查任务详情"""
    service = AuditService(db)
    task = service.get_task_by_id(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return task_to_dict(task)


@router.post("/task/{task_id}/start")
def start_task(
    task_id: int,
    request: StartTaskRequest = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """启动审查任务"""
    service = AuditService(db)
    config_id = request.config_id if request else None
    task = service.start_task(task_id, config_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"success": True, "message": "审查任务已启动"}


@router.post("/task/{task_id}/pause")
def pause_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """暂停审查任务"""
    service = AuditService(db)
    task = service.pause_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"success": True, "message": "审查任务已暂停"}


@router.post("/task/{task_id}/cancel")
def cancel_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """取消审查任务"""
    service = AuditService(db)
    task = service.cancel_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"success": True, "message": "审查任务已取消"}


@router.get("/task/{task_id}/stream")
async def stream_audit_result(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """流式获取审查结果 - 执行真正的审查逻辑"""
    service = AuditService(db)
    task = service.get_task_by_id(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    async def generate():
        try:
            yield f"data: {json.dumps({'content': '正在解析文档...', 'is_end': False}, ensure_ascii=False)}\n\n"
            service.update_task_status(task_id, "parsing", 10)
            await asyncio.sleep(0.3)
            
            file_path = task.document_path
            if not os.path.exists(file_path):
                yield f"data: {json.dumps({'content': f'文件不存在: {file_path}', 'is_end': True}, ensure_ascii=False)}\n\n"
                service.update_task_status(task_id, "failed", 0)
                return
            
            file_content = None
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                from utils.pdf_parser import PDFParser
                parser = PDFParser()
                file_content = parser.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(file_path)
                file_content = "\n".join([para.text for para in doc.paragraphs])
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
            else:
                yield f"data: {json.dumps({'content': f'不支持的文件格式: {file_ext}', 'is_end': True}, ensure_ascii=False)}\n\n"
                service.update_task_status(task_id, "failed", 0)
                return
            
            if not file_content:
                yield f"data: {json.dumps({'content': '无法提取文件内容', 'is_end': True}, ensure_ascii=False)}\n\n"
                service.update_task_status(task_id, "failed", 0)
                return
            
            yield f"data: {json.dumps({'content': f'文档解析完成，共 {len(file_content)} 字符', 'is_end': False}, ensure_ascii=False)}\n\n"
            service.update_task_status(task_id, "analyzing", 30)
            await asyncio.sleep(0.3)
            
            config = None
            if task.config_id:
                config = service.get_config_by_id(task.config_id)
            
            audit_dimensions = config.audit_dimensions if config and config.audit_dimensions else ["compliance", "consistency", "format"]
            focus_keywords = config.focus_keywords if config and config.focus_keywords else []
            
            dimension_names = {
                "compliance": "合规性",
                "consistency": "一致性",
                "format": "形式审查"
            }
            
            question_parts = [f"请对以下制度文件进行全面审查，文件名：{task.document_name}"]
            
            if "compliance" in audit_dimensions:
                question_parts.append("\n\n【合规性审查】请检查：\n1. 是否违反现行法律法规\n2. 是否与上位制度冲突\n3. 条款是否存在法律风险")
            
            if "consistency" in audit_dimensions:
                question_parts.append("\n\n【一致性审查】请检查：\n1. 条款之间是否一致\n2. 引用是否准确\n3. 术语使用是否统一")
            
            if "format" in audit_dimensions:
                question_parts.append("\n\n【形式审查】请检查：\n1. 是否有错别字\n2. 格式是否规范\n3. 逻辑是否清晰")
            
            if focus_keywords:
                question_parts.append(f"\n\n【重点关注】请特别关注以下关键词相关内容：{', '.join(focus_keywords)}")
            
            question_parts.append("\n\n请按照以下格式输出审查结果：\n")
            question_parts.append("【问题序号】问题类型：合规性/一致性/形式\n")
            question_parts.append("位置：具体位置\n")
            question_parts.append("原文：相关原文内容\n")
            question_parts.append("问题描述：具体问题描述\n")
            question_parts.append("风险等级：高/中/低\n")
            question_parts.append("修改建议：具体修改建议\n")
            
            question = "".join(question_parts)
            
            yield f"data: {json.dumps({'content': '正在调用AI进行审查...', 'is_end': False}, ensure_ascii=False)}\n\n"
            service.update_task_status(task_id, "analyzing", 50)
            
            agent = create_audit_agent(enable_tools=True)
            
            loop = asyncio.get_event_loop()
            agent_result = await loop.run_in_executor(
                None,
                lambda: agent.chat_with_tools(
                    question=question,
                    messages=[],
                    file_content=file_content,
                    file_paths=[file_path],
                    conversation_id=task.conversation_id
                )
            )
            
            response_content = agent_result.get("response", "")
            
            yield f"data: {json.dumps({'content': '正在解析审查结果...', 'is_end': False}, ensure_ascii=False)}\n\n"
            service.update_task_status(task_id, "analyzing", 80)
            await asyncio.sleep(0.3)
            
            # 先检查是否已存在结果
            existing_result = service.get_result_by_task_id(task_id)
            if existing_result:
                result = existing_result
            else:
                result = service.create_result(
                    task_id=task_id,
                    document_name=task.document_name
                )
            
            issues = parse_audit_issues(response_content, result.id)
            
            for issue_data in issues:
                issue = service.create_issue(
                    result_id=result.id,
                    issue_type=issue_data.get("issue_type", "compliance"),
                    severity=issue_data.get("severity", "medium"),
                    location=issue_data.get("location", ""),
                    original_text=issue_data.get("original_text", ""),
                    issue_description=issue_data.get("issue_description", ""),
                    legal_basis=issue_data.get("legal_basis", ""),
                    suggestion=issue_data.get("suggestion", "")
                )
                
                yield f"data: {json.dumps({'content': f'发现问题：{issue.issue_description[:100]}...', 'is_end': False}, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.1)
            
            service.update_task_status(task_id, "completed", 100)
            
            yield f"data: {json.dumps({'content': f'审查完成，共发现 {len(issues)} 个问题', 'is_end': True, 'result_id': result.id}, ensure_ascii=False)}\n\n"
            
        except Exception as e:
            logger.error(f"审查任务执行失败: {str(e)}", exc_info=True)
            service.update_task_status(task_id, "failed", 0)
            yield f"data: {json.dumps({'content': f'审查失败: {str(e)}', 'is_end': True}, ensure_ascii=False)}\n\n"
    
    return StreamingResponse(generate(), media_type="text/event-stream")


def parse_audit_issues(content: str, result_id: int) -> List[dict]:
    """解析AI返回的审查结果，提取问题列表"""
    issues = []
    
    issue_pattern = r'【问题(\d+)】.*?(?=【问题\d+】|$)'
    issue_blocks = re.findall(issue_pattern, content, re.DOTALL)
    
    if not issue_blocks:
        blocks = re.split(r'\n(?=\d+[\.、）])', content)
        for block in blocks:
            if block.strip():
                issue = parse_single_issue(block, result_id)
                if issue:
                    issues.append(issue)
    else:
        for match in re.finditer(issue_pattern, content, re.DOTALL):
            block = match.group(0)
            issue = parse_single_issue(block, result_id)
            if issue:
                issues.append(issue)
    
    return issues


def parse_single_issue(block: str, result_id: int) -> Optional[dict]:
    """解析单个问题块"""
    issue = {
        "result_id": result_id,
        "issue_type": "compliance",
        "severity": "medium",
        "location": "",
        "original_text": "",
        "issue_description": "",
        "legal_basis": "",
        "suggestion": ""
    }
    
    type_match = re.search(r'问题类型[：:]\s*(合规性|一致性|形式)', block)
    if type_match:
        type_map = {"合规性": "compliance", "一致性": "consistency", "形式": "format"}
        issue["issue_type"] = type_map.get(type_match.group(1), "compliance")
    
    location_match = re.search(r'位置[：:]\s*(.+?)(?:\n|$)', block)
    if location_match:
        issue["location"] = location_match.group(1).strip()
    
    original_match = re.search(r'原文[：:]\s*(.+?)(?=\n(?:问题|风险|修改|建议)|$)', block, re.DOTALL)
    if original_match:
        issue["original_text"] = original_match.group(1).strip()
    
    desc_match = re.search(r'问题描述[：:]\s*(.+?)(?=\n(?:风险|修改|建议|原文)|$)', block, re.DOTALL)
    if desc_match:
        issue["issue_description"] = desc_match.group(1).strip()
    
    severity_match = re.search(r'风险等级[：:]\s*(高|中|低)', block)
    if severity_match:
        severity_map = {"高": "high", "中": "medium", "低": "low"}
        issue["severity"] = severity_map.get(severity_match.group(1), "medium")
    
    suggestion_match = re.search(r'修改建议[：:]\s*(.+?)(?=\n(?:【问题|\d+[\.、）])|$)', block, re.DOTALL)
    if suggestion_match:
        issue["suggestion"] = suggestion_match.group(1).strip()
    
    if issue["issue_description"]:
        return issue
    
    return None


@router.get("/task/{task_id}/trails")
def get_task_trails(task_id: int, db: Session = Depends(get_db)):
    """获取审查轨迹"""
    service = AuditService(db)
    result = service.get_trails_by_task_id(task_id)
    return {
        "total": result["total"],
        "items": [trail_to_dict(t) for t in result["items"]]
    }


@router.post("/task/{task_id}/trail/export")
def export_trail(task_id: int, db: Session = Depends(get_db)):
    """导出审查轨迹"""
    service = AuditService(db)
    task = service.get_task_by_id(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    return {
        "download_url": f"/downloads/trail_{task_id}.xlsx",
        "file_name": f"{task.document_name}_审查轨迹.xlsx"
    }


# ==================== 审查结果接口 ====================

@router.get("/result/{result_id}")
def get_result(result_id: int, db: Session = Depends(get_db)):
    """获取审查结果详情"""
    service = AuditService(db)
    result = service.get_result_by_id(result_id)
    if not result:
        raise HTTPException(status_code=404, detail="结果不存在")
    return result_to_dict(result)


@router.get("/result/{result_id}/issues")
def get_result_issues(
    result_id: int,
    issue_type: Optional[str] = Query(None),
    severity: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    """获取审查问题列表"""
    service = AuditService(db)
    result = service.get_issues_by_result_id(result_id, issue_type, severity)
    return {
        "total": result["total"],
        "items": [issue_to_dict(i) for i in result["items"]]
    }


@router.get("/result/{result_id}/document")
def get_document_content(result_id: int, db: Session = Depends(get_db)):
    """获取文档内容"""
    service = AuditService(db)
    result = service.get_result_by_id(result_id)
    if not result:
        raise HTTPException(status_code=404, detail="结果不存在")
    
    task = service.get_task_by_id(result.task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    try:
        file_path = task.document_path
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            from utils.pdf_parser import PDFParser
            parser = PDFParser()
            content = parser.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {file_ext}")
        
        return {"content": content, "file_type": file_ext}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"读取文档失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"读取文档失败: {str(e)}")


@router.post("/result/{result_id}/export")
def export_report(result_id: int, request: ExportReportRequest, db: Session = Depends(get_db)):
    """导出审查报告"""
    from datetime import datetime
    
    service = AuditService(db)
    result = service.get_result_by_id(result_id)
    if not result:
        raise HTTPException(status_code=404, detail="结果不存在")
    
    task = service.get_task_by_id(result.task_id)
    issues_result = service.get_issues_by_result_id(result_id)
    issues = issues_result["items"]
    
    reports_dir = "uploads/reports"
    os.makedirs(reports_dir, exist_ok=True)
    
    document_name = result.document_name or "未命名文档"
    name_without_ext = os.path.splitext(document_name)[0]
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    
    if request.format == "word":
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        title = doc.add_heading(f"{name_without_ext} 审查报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"审查时间：{result.created_at.strftime('%Y-%m-%d %H:%M:%S') if result.created_at else ''}")
        doc.add_paragraph(f"风险等级：{result.risk_level.value if result.risk_level else '未评估'}")
        doc.add_paragraph(f"问题总数：{result.total_issues}")
        doc.add_paragraph(f"合规性问题：{result.compliance_issues}")
        doc.add_paragraph(f"一致性问题：{result.consistency_issues}")
        doc.add_paragraph(f"形式问题：{result.format_issues}")
        
        doc.add_heading("审查问题详情", level=1)
        
        for i, issue in enumerate(issues, 1):
            doc.add_heading(f"问题 {i}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("问题类型：").bold = True
            p.add_run(issue.issue_type.value if issue.issue_type else "")
            
            p = doc.add_paragraph()
            p.add_run("风险等级：").bold = True
            p.add_run(issue.severity.value if issue.severity else "")
            
            if issue.location:
                p = doc.add_paragraph()
                p.add_run("位置：").bold = True
                p.add_run(issue.location)
            
            if issue.original_text:
                p = doc.add_paragraph()
                p.add_run("原文内容：").bold = True
                doc.add_paragraph(issue.original_text)
            
            if issue.issue_description:
                p = doc.add_paragraph()
                p.add_run("问题描述：").bold = True
                doc.add_paragraph(issue.issue_description)
            
            if issue.legal_basis:
                p = doc.add_paragraph()
                p.add_run("法律依据：").bold = True
                doc.add_paragraph(issue.legal_basis)
            
            if issue.suggestion:
                p = doc.add_paragraph()
                p.add_run("修改建议：").bold = True
                doc.add_paragraph(issue.suggestion)
        
        file_name = f"{name_without_ext}_审查报告_{timestamp}.docx"
        file_path = os.path.join(reports_dir, file_name)
        doc.save(file_path)
        
        return {
            "download_url": f"/api/file/download/{file_name}",
            "file_name": file_name
        }
    else:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import cm
        
        chinese_font = None
        font_paths = [
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simsun.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if font_path.endswith('.ttc'):
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                    else:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    chinese_font = 'ChineseFont'
                    logger.info(f"成功加载中文字体: {font_path}")
                    break
                except Exception as e:
                    logger.warning(f"加载字体失败 {font_path}: {e}")
                    continue
        
        if not chinese_font:
            logger.warning("未找到中文字体，使用默认字体，中文可能无法正常显示")
            chinese_font = 'Helvetica'
        
        file_name = f"{name_without_ext}_审查报告_{timestamp}.pdf"
        file_path = os.path.join(reports_dir, file_name)
        
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=18,
            alignment=1,
            leading=24
        )
        
        heading_style = ParagraphStyle(
            'ChineseHeading',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=14,
            leading=18
        )
        
        normal_style = ParagraphStyle(
            'ChineseNormal',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=14
        )
        
        story = []
        
        story.append(Paragraph(f"{name_without_ext} 审查报告", title_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph(f"审查时间：{result.created_at.strftime('%Y-%m-%d %H:%M:%S') if result.created_at else ''}", normal_style))
        story.append(Paragraph(f"风险等级：{result.risk_level.value if result.risk_level else '未评估'}", normal_style))
        story.append(Paragraph(f"问题总数：{result.total_issues}", normal_style))
        story.append(Paragraph(f"合规性问题：{result.compliance_issues}", normal_style))
        story.append(Paragraph(f"一致性问题：{result.consistency_issues}", normal_style))
        story.append(Paragraph(f"形式问题：{result.format_issues}", normal_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("审查问题详情", heading_style))
        story.append(Spacer(1, 10))
        
        for i, issue in enumerate(issues, 1):
            story.append(Paragraph(f"问题 {i}", heading_style))
            
            story.append(Paragraph(f"问题类型：{issue.issue_type.value if issue.issue_type else ''}", normal_style))
            story.append(Paragraph(f"风险等级：{issue.severity.value if issue.severity else ''}", normal_style))
            
            if issue.location:
                story.append(Paragraph(f"位置：{issue.location}", normal_style))
            
            if issue.original_text:
                story.append(Paragraph(f"原文内容：{issue.original_text}", normal_style))
            
            if issue.issue_description:
                story.append(Paragraph(f"问题描述：{issue.issue_description}", normal_style))
            
            if issue.legal_basis:
                story.append(Paragraph(f"法律依据：{issue.legal_basis}", normal_style))
            
            if issue.suggestion:
                story.append(Paragraph(f"修改建议：{issue.suggestion}", normal_style))
            
            story.append(Spacer(1, 10))
        
        doc.build(story)
        
        return {
            "download_url": f"/api/file/download/{file_name}",
            "file_name": file_name
        }


# ==================== 问题管理接口 ====================

@router.put("/issue/{issue_id}/status")
def update_issue_status(
    issue_id: int,
    request: UpdateIssueStatusRequest,
    db: Session = Depends(get_db)
):
    """更新问题状态"""
    service = AuditService(db)
    issue = service.update_issue_status(
        issue_id=issue_id,
        status=request.status,
        suggestion=request.suggestion,
        reject_reason=request.reject_reason
    )
    if not issue:
        raise HTTPException(status_code=404, detail="问题不存在")
    return {"success": True, "message": "问题状态已更新"}


# ==================== 配置管理接口 ====================

@router.post("/config/create")
def create_config(request: CreateConfigRequest, db: Session = Depends(get_db)):
    """创建审查配置"""
    service = AuditService(db)
    config = service.create_config(
        name=request.name,
        audit_dimensions=request.audit_dimensions,
        focus_keywords=request.focus_keywords,
        checklist_ids=request.checklist_ids,
        is_default=request.is_default
    )
    return config_to_dict(config)


@router.get("/configs")
def get_configs(db: Session = Depends(get_db)):
    """获取所有审查配置"""
    service = AuditService(db)
    configs = service.get_configs()
    return [config_to_dict(c) for c in configs]


# ==================== 清单管理接口 ====================

@router.get("/checklists")
def get_checklists(db: Session = Depends(get_db)):
    """获取审查清单列表"""
    service = AuditService(db)
    result = service.get_checklists(is_active=True)
    return {
        "total": result["total"],
        "items": [checklist_to_dict(c) for c in result["items"]]
    }


# ==================== 版本比对接口 ====================

@router.post("/version-compare/create")
def create_version_compare(request: CreateVersionCompareRequest, db: Session = Depends(get_db)):
    """创建版本比对任务"""
    service = AuditService(db)
    task = service.create_version_compare_task(
        old_document_path=request.old_document_path,
        new_document_path=request.new_document_path,
        config_id=request.config_id
    )
    return {"id": task.id}


@router.get("/version-compare/{task_id}")
def get_version_compare_result(task_id: int, db: Session = Depends(get_db)):
    """获取版本比对结果"""
    service = AuditService(db)
    task = service.get_version_compare_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return version_compare_to_dict(task)


# ==================== 审核确认接口 ====================

@router.post("/result/{result_id}/start-review")
def start_review(
    result_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """开始审核 - 将审查结果状态改为审核中"""
    service = AuditService(db)
    result = service.start_review(result_id, current_user.id)
    if not result:
        raise HTTPException(status_code=400, detail="无法开始审核，结果不存在或状态不正确")
    return {"success": True, "message": "已开始审核", "result": result_to_dict(result)}


@router.post("/result/{result_id}/confirm")
def confirm_result(
    result_id: int,
    request: ConfirmResultRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """确认审查结果 - 完成审核"""
    service = AuditService(db)
    result = service.confirm_result(
        result_id=result_id,
        reviewer_id=current_user.id,
        comment=request.comment
    )
    if not result:
        raise HTTPException(status_code=404, detail="结果不存在")
    return {"success": True, "message": "审查结果已确认", "result": result_to_dict(result)}


@router.post("/result/{result_id}/reject")
def reject_result(
    result_id: int,
    request: RejectResultRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """驳回审查结果 - 需要重新审查"""
    service = AuditService(db)
    result = service.reject_result(
        result_id=result_id,
        reviewer_id=current_user.id,
        reason=request.reason
    )
    if not result:
        raise HTTPException(status_code=404, detail="结果不存在")
    return {"success": True, "message": "审查结果已驳回", "result": result_to_dict(result)}


@router.get("/result/{result_id}/review-statistics")
def get_review_statistics(
    result_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取审核统计信息"""
    service = AuditService(db)
    stats = service.get_review_statistics(result_id)
    if not stats:
        raise HTTPException(status_code=404, detail="结果不存在")
    return stats


@router.post("/result/{result_id}/issues/batch")
def batch_update_issues(
    result_id: int,
    request: BatchUpdateIssuesRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """批量更新问题状态"""
    service = AuditService(db)
    result = service.batch_update_issues(
        result_id=result_id,
        issue_ids=request.issue_ids,
        status=request.status,
        user_id=current_user.id
    )
    if not result.get("success"):
        raise HTTPException(status_code=400, detail=result.get("message", "批量更新失败"))
    return result


@router.post("/result/{result_id}/issues/accept-all")
def accept_all_issues(
    result_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """接受所有问题"""
    service = AuditService(db)
    issues_result = service.get_issues_by_result_id(result_id)
    issue_ids = [issue.id for issue in issues_result["items"]]
    
    if not issue_ids:
        return {"success": True, "updated_count": 0, "message": "没有需要处理的问题"}
    
    result = service.batch_update_issues(
        result_id=result_id,
        issue_ids=issue_ids,
        status="accepted",
        user_id=current_user.id
    )
    return result


@router.post("/issue/{issue_id}/accept")
def accept_issue(
    issue_id: int,
    suggestion: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """接受单个问题"""
    service = AuditService(db)
    issue = service.update_issue_status(issue_id, "accepted", suggestion=suggestion)
    if not issue:
        raise HTTPException(status_code=404, detail="问题不存在")
    return {"success": True, "message": "问题已接受", "issue": issue_to_dict(issue)}


@router.post("/issue/{issue_id}/reject")
def reject_issue(
    issue_id: int,
    reject_reason: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """拒绝单个问题"""
    service = AuditService(db)
    issue = service.update_issue_status(issue_id, "rejected", reject_reason=reject_reason)
    if not issue:
        raise HTTPException(status_code=404, detail="问题不存在")
    return {"success": True, "message": "问题已拒绝", "issue": issue_to_dict(issue)}
