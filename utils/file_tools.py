"""
文件操作工具模块
提供文件读取、修改、标注等功能
"""

import os
import re
import json
import tempfile
from typing import Optional, List, Dict, Any
from pathlib import Path
from urllib.parse import urlparse, unquote
import httpx
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import red, black
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from config import settings


class FileOperationTools:
    """
    文件操作工具类
    
    提供文件读取、修改、标注等功能，支持PDF和Word文档
    
    Attributes:
        upload_dir: 上传文件目录
        output_dir: 输出文件目录
        converted_dir: 转换文件目录
    """
    
    def __init__(self):
        """
        初始化文件操作工具
        """
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.output_dir = self.upload_dir / "modified"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.converted_dir = self.upload_dir / "converted"
        self.converted_dir.mkdir(parents=True, exist_ok=True)
        self._init_fonts()
    
    def _init_fonts(self):
        """
        初始化中文字体支持
        """
        try:
            font_paths = [
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/msyh.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            ]
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    break
        except Exception:
            pass
    
    def read_file_content(self, file_path: str) -> Dict[str, Any]:
        """
        读取文件内容
        
        Args:
            file_path: 文件路径
        
        Returns:
            Dict[str, Any]: 包含文件内容和元数据的字典
        """
        if not os.path.exists(file_path):
            return {"success": False, "error": "文件不存在"}
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == ".pdf":
                return self._read_pdf(file_path)
            elif file_ext in [".docx", ".doc"]:
                return self._read_docx(file_path)
            elif file_ext == ".txt":
                return self._read_txt(file_path)
            else:
                return {"success": False, "error": f"不支持的文件类型: {file_ext}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        读取PDF文件内容
        
        Args:
            file_path: PDF文件路径
        
        Returns:
            Dict[str, Any]: 文件内容
        """
        reader = PdfReader(file_path)
        pages_content = []
        
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            pages_content.append({
                "page": i,
                "content": text
            })
        
        return {
            "success": True,
            "file_type": "pdf",
            "page_count": len(reader.pages),
            "pages": pages_content,
            "full_content": "\n\n".join([p["content"] for p in pages_content])
        }
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """
        读取Word文件内容
        
        Args:
            file_path: Word文件路径
        
        Returns:
            Dict[str, Any]: 文件内容
        """
        doc = Document(file_path)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs, start=1):
            paragraphs.append({
                "index": i,
                "content": para.text
            })
        
        return {
            "success": True,
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "paragraphs": paragraphs,
            "full_content": "\n".join([p["content"] for p in paragraphs])
        }
    
    def _read_txt(self, file_path: str) -> Dict[str, Any]:
        """
        读取文本文件内容
        
        Args:
            file_path: 文本文件路径
        
        Returns:
            Dict[str, Any]: 文件内容
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "success": True,
            "file_type": "txt",
            "full_content": content
        }
    
    def highlight_text_in_pdf(
        self,
        file_path: str,
        highlight_texts: List[str],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        在PDF中标注文字（生成新的PDF，标注的文字用红色显示）
        
        Args:
            file_path: 原PDF文件路径
            highlight_texts: 需要标注的文字列表
            output_filename: 输出文件名
        
        Returns:
            Dict[str, Any]: 操作结果
        """
        try:
            content_result = self._read_pdf(file_path)
            if not content_result.get("success"):
                return content_result
            
            full_content = content_result["full_content"]
            
            output_filename = output_filename or f"highlighted_{Path(file_path).stem}.pdf"
            output_path = self.output_dir / output_filename
            
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4
            )
            
            styles = getSampleStyleSheet()
            try:
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=10,
                    leading=14
                )
            except Exception:
                normal_style = styles['Normal']
            
            story = []
            
            for text in highlight_texts:
                full_content = self._highlight_text(full_content, text)
            
            paragraphs = full_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    try:
                        story.append(Paragraph(para_text, normal_style))
                    except Exception:
                        story.append(Paragraph(para_text.encode('utf-8', errors='ignore').decode('utf-8'), normal_style))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            
            return {
                "success": True,
                "output_path": str(output_path),
                "output_filename": output_filename,
                "highlighted_texts": highlight_texts,
                "message": f"已生成标注后的PDF文件: {output_filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _highlight_text(self, content: str, text_to_highlight: str) -> str:
        """
        将内容中指定文字标记为红色（使用XML标记）
        
        Args:
            content: 原始内容
            text_to_highlight: 需要标注的文字
        
        Returns:
            str: 标注后的内容
        """
        pattern = re.compile(re.escape(text_to_highlight), re.IGNORECASE)
        return pattern.sub(f'<font color="red">{text_to_highlight}</font>', content)
    
    def modify_text_in_pdf(
        self,
        file_path: str,
        modifications: List[Dict[str, str]],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        修改PDF中的文字内容
        
        Args:
            file_path: 原PDF文件路径
            modifications: 修改列表，每项包含 {"old_text": "旧文字", "new_text": "新文字"}
            output_filename: 输出文件名
        
        Returns:
            Dict[str, Any]: 操作结果
        """
        try:
            content_result = self._read_pdf(file_path)
            if not content_result.get("success"):
                return content_result
            
            full_content = content_result["full_content"]
            modification_log = []
            
            for mod in modifications:
                old_text = mod.get("old_text", "")
                new_text = mod.get("new_text", "")
                
                if old_text and old_text in full_content:
                    full_content = full_content.replace(old_text, new_text)
                    modification_log.append({
                        "old_text": old_text,
                        "new_text": new_text,
                        "status": "success"
                    })
                else:
                    modification_log.append({
                        "old_text": old_text,
                        "new_text": new_text,
                        "status": "not_found"
                    })
            
            output_filename = output_filename or f"modified_{Path(file_path).stem}.pdf"
            output_path = self.output_dir / output_filename
            
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4
            )
            
            styles = getSampleStyleSheet()
            try:
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=10,
                    leading=14
                )
            except Exception:
                normal_style = styles['Normal']
            
            story = []
            paragraphs = full_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    try:
                        story.append(Paragraph(para_text, normal_style))
                    except Exception:
                        story.append(Paragraph(para_text.encode('utf-8', errors='ignore').decode('utf-8'), normal_style))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            
            return {
                "success": True,
                "output_path": str(output_path),
                "output_filename": output_filename,
                "modifications": modification_log,
                "message": f"已生成修改后的PDF文件: {output_filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def create_highlighted_docx(
        self,
        file_path: str,
        highlight_texts: List[str],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        创建标注后的Word文档
        
        Args:
            file_path: 原文件路径
            highlight_texts: 需要标注的文字列表
            output_filename: 输出文件名
        
        Returns:
            Dict[str, Any]: 操作结果
        """
        try:
            content_result = self.read_file_content(file_path)
            if not content_result.get("success"):
                return content_result
            
            full_content = content_result["full_content"]
            
            output_filename = output_filename or f"highlighted_{Path(file_path).stem}.docx"
            output_path = self.output_dir / output_filename
            
            doc = Document()
            
            paragraphs = full_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    self._add_highlighted_paragraph(para, para_text, highlight_texts)
            
            doc.save(str(output_path))
            
            return {
                "success": True,
                "output_path": str(output_path),
                "output_filename": output_filename,
                "highlighted_texts": highlight_texts,
                "message": f"已生成标注后的Word文件: {output_filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _add_highlighted_paragraph(self, para, text: str, highlight_texts: List[str]):
        """
        添加带标注的段落
        
        Args:
            para: 段落对象
            text: 段落文本
            highlight_texts: 需要标注的文字列表
        """
        if not highlight_texts:
            para.add_run(text)
            return
        
        pattern = '|'.join(re.escape(t) for t in highlight_texts)
        parts = re.split(f'({pattern})', text)
        
        for part in parts:
            if part in highlight_texts:
                run = para.add_run(part)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True
            else:
                para.add_run(part)
    
    def modify_text_in_docx(
        self,
        file_path: str,
        modifications: List[Dict[str, str]],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        修改Word文档中的文字内容
        
        Args:
            file_path: 原Word文件路径
            modifications: 修改列表
            output_filename: 输出文件名
        
        Returns:
            Dict[str, Any]: 操作结果
        """
        try:
            doc = Document(file_path)
            modification_log = []
            
            for mod in modifications:
                old_text = mod.get("old_text", "")
                new_text = mod.get("new_text", "")
                
                found = False
                for para in doc.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                found = True
                
                modification_log.append({
                    "old_text": old_text,
                    "new_text": new_text,
                    "status": "success" if found else "not_found"
                })
            
            output_filename = output_filename or f"modified_{Path(file_path).stem}.docx"
            output_path = self.output_dir / output_filename
            
            doc.save(str(output_path))
            
            return {
                "success": True,
                "output_path": str(output_path),
                "output_filename": output_filename,
                "modifications": modification_log,
                "message": f"已生成修改后的Word文件: {output_filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def add_review_comments(
        self,
        file_path: str,
        comments: List[Dict[str, Any]],
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        添加审查意见到文档
        
        Args:
            file_path: 原文件路径
            comments: 审查意见列表，每项包含 {"text": "原文", "comment": "意见", "risk_level": "风险等级"}
            output_filename: 输出文件名
        
        Returns:
            Dict[str, Any]: 操作结果
        """
        try:
            content_result = self.read_file_content(file_path)
            if not content_result.get("success"):
                return content_result
            
            output_filename = output_filename or f"reviewed_{Path(file_path).stem}.docx"
            output_path = self.output_dir / output_filename
            
            doc = Document()
            
            doc.add_heading('制度审查报告', 0)
            doc.add_paragraph('')
            
            doc.add_heading('一、审查意见汇总', level=1)
            
            for i, comment in enumerate(comments, start=1):
                risk_level = comment.get("risk_level", "中")
                original_text = comment.get("text", "")
                review_comment = comment.get("comment", "")
                
                if risk_level == "高":
                    risk_color = RGBColor(255, 0, 0)
                elif risk_level == "中":
                    risk_color = RGBColor(255, 165, 0)
                else:
                    risk_color = RGBColor(0, 128, 0)
                
                doc.add_heading(f'审查点 {i}', level=2)
                
                para = doc.add_paragraph()
                run = para.add_run(f'风险等级: {risk_level}')
                run.font.color.rgb = risk_color
                run.font.bold = True
                
                doc.add_paragraph(f'原文内容: {original_text}')
                doc.add_paragraph(f'审查意见: {review_comment}')
                doc.add_paragraph('')
            
            doc.add_heading('二、原文内容（标注版）', level=1)
            
            highlight_texts = [c.get("text", "") for c in comments if c.get("text")]
            full_content = content_result["full_content"]
            
            paragraphs = full_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    self._add_highlighted_paragraph(para, para_text, highlight_texts)
            
            doc.save(str(output_path))
            
            return {
                "success": True,
                "output_path": str(output_path),
                "output_filename": output_filename,
                "comments_count": len(comments),
                "message": f"已生成审查报告: {output_filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def get_output_file(self, filename: str) -> Optional[str]:
        """
        获取输出文件路径
        
        Args:
            filename: 文件名
        
        Returns:
            Optional[str]: 文件路径，不存在则返回None
        """
        file_path = self.output_dir / filename
        if file_path.exists():
            return str(file_path)
        
        converted_file = self.converted_dir / filename
        if converted_file.exists():
            return str(converted_file)
        
        return None
    
    def convert_to_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        将文件转换为PDF格式
        
        支持 Word(.docx)、TXT 等格式转换为 PDF
        支持本地文件路径和URL
        
        Args:
            file_path: 原文件路径（相对路径、绝对路径或URL）
        
        Returns:
            Dict[str, Any]: 转换结果，包含 success、download_url 或 error
        """
        temp_file = None
        
        try:
            parsed_url = urlparse(file_path)
            is_url = parsed_url.scheme in ['http', 'https']
            
            if is_url:
                if self._is_local_url(file_path):
                    local_file_path = self._get_local_file_path_from_url(file_path)
                    if not local_file_path or not os.path.exists(local_file_path):
                        return {"success": False, "error": "文件不存在"}
                    filename = unquote(Path(parsed_url.path).name)
                    file_ext = Path(filename).suffix.lower()
                else:
                    result = self._download_file_from_url(file_path)
                    if not result.get("success"):
                        return result
                    local_file_path = result["file_path"]
                    temp_file = local_file_path
                    filename = unquote(Path(parsed_url.path).name)
                    file_ext = Path(filename).suffix.lower()
            else:
                if not os.path.exists(file_path):
                    return {"success": False, "error": "文件不存在"}
                
                local_file_path = file_path
                file_ext = Path(file_path).suffix.lower()
                filename = Path(file_path).name
            
            if file_ext == ".pdf":
                return {"success": False, "error": "文件已经是PDF格式，无需转换"}
            
            if file_ext not in [".docx", ".doc", ".txt"]:
                return {"success": False, "error": f"不支持的文件类型: {file_ext}，仅支持 Word(.docx) 和 TXT 格式"}
            
            output_filename = f"{Path(filename).stem}.pdf"
            output_path = self.converted_dir / output_filename
            
            if file_ext in [".docx", ".doc"]:
                result = self._convert_docx_to_pdf(local_file_path, str(output_path))
            elif file_ext == ".txt":
                result = self._convert_txt_to_pdf(local_file_path, str(output_path))
            else:
                return {"success": False, "error": f"不支持的文件类型: {file_ext}"}
            
            if result.get("success"):
                download_url = f"uploads/converted/{output_filename}"
                return {
                    "success": True,
                    "download_url": download_url,
                    "output_path": str(output_path),
                    "output_filename": output_filename,
                    "message": f"文件转换成功: {output_filename}"
                }
            else:
                return result
                
        except Exception as e:
            return {"success": False, "error": f"转换失败: {str(e)}"}
        finally:
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception:
                    pass
    
    def _is_local_url(self, url: str) -> bool:
        """
        判断URL是否为本地URL
        
        Args:
            url: URL字符串
        
        Returns:
            bool: 是否为本地URL
        """
        try:
            parsed = urlparse(url)
            return parsed.scheme in ['http', 'https'] and parsed.hostname in ['localhost', '127.0.0.1']
        except Exception:
            return False
    
    def _get_local_file_path_from_url(self, url: str) -> Optional[str]:
        """
        从本地URL提取本地文件路径
        
        Args:
            url: 本地URL
        
        Returns:
            Optional[str]: 本地文件路径，如果无法提取则返回None
        """
        try:
            parsed = urlparse(url)
            path = unquote(parsed.path)
            
            if path.startswith('/api/file/download/'):
                filename = path.replace('/api/file/download/', '')
                
                output_file = self.get_output_file(filename)
                if output_file:
                    return output_file
                
                converted_file = self.converted_dir / filename
                if converted_file.exists():
                    return str(converted_file)
                
                temp_file = self.upload_dir / "temp" / filename
                if temp_file.exists():
                    return str(temp_file)
                
                kb_file = self.upload_dir / "knowledge_base" / filename
                if kb_file.exists():
                    return str(kb_file)
                
                direct_file = self.upload_dir / filename
                if direct_file.exists():
                    return str(direct_file)
            
            return None
            
        except Exception:
            return None
    
    def _download_file_from_url(self, url: str) -> Dict[str, Any]:
        """
        从URL下载文件到临时目录
        
        Args:
            url: 文件URL
        
        Returns:
            Dict[str, Any]: 下载结果，包含 success 和 file_path 或 error
        """
        try:
            response = httpx.get(url, follow_redirects=True, timeout=30.0)
            
            if response.status_code != 200:
                return {"success": False, "error": f"下载文件失败: HTTP {response.status_code}"}
            
            parsed_url = urlparse(url)
            filename = unquote(Path(parsed_url.path).name)
            
            if not filename:
                filename = "downloaded_file"
            
            temp_dir = tempfile.gettempdir()
            temp_file_path = os.path.join(temp_dir, filename)
            
            with open(temp_file_path, 'wb') as f:
                f.write(response.content)
            
            return {
                "success": True,
                "file_path": temp_file_path,
                "filename": filename
            }
            
        except Exception as e:
            return {"success": False, "error": f"下载文件失败: {str(e)}"}
    
    def _convert_docx_to_pdf(self, file_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 Word 文档转换为 PDF
        
        Args:
            file_path: Word 文件路径
            output_path: 输出 PDF 文件路径
        
        Returns:
            Dict[str, Any]: 转换结果
        """
        try:
            doc = Document(file_path)
            
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4
            )
            
            styles = getSampleStyleSheet()
            try:
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=10,
                    leading=14
                )
            except Exception:
                normal_style = styles['Normal']
            
            story = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    try:
                        story.append(Paragraph(para.text, normal_style))
                    except Exception:
                        story.append(Paragraph(para.text.encode('utf-8', errors='ignore').decode('utf-8'), normal_style))
                    story.append(Spacer(1, 6))
            
            pdf_doc.build(story)
            
            return {"success": True}
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _convert_txt_to_pdf(self, file_path: str, output_path: str) -> Dict[str, Any]:
        """
        将 TXT 文件转换为 PDF
        
        Args:
            file_path: TXT 文件路径
            output_path: 输出 PDF 文件路径
        
        Returns:
            Dict[str, Any]: 转换结果
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4
            )
            
            styles = getSampleStyleSheet()
            try:
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=10,
                    leading=14
                )
            except Exception:
                normal_style = styles['Normal']
            
            story = []
            
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    try:
                        story.append(Paragraph(para_text, normal_style))
                    except Exception:
                        story.append(Paragraph(para_text.encode('utf-8', errors='ignore').decode('utf-8'), normal_style))
                    story.append(Spacer(1, 6))
            
            pdf_doc.build(story)
            
            return {"success": True}
            
        except Exception as e:
            return {"success": False, "error": str(e)}


file_tools = FileOperationTools()
